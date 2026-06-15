from langchain.tools import tool
import pdfplumber
from docx import Document


@tool
def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text from a PDF file (works for text-selectable PDFs).
    For scanned PDFs, use OCR tools instead.

    Args:
        pdf_path: Path to the PDF file

    Returns:
        Extracted text from all pages
    """
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += f"--- Page {page_num} ---\n{page_text}\n"
        return text if text else "No text found in PDF"
    except Exception as e:
        return f"Error reading PDF {pdf_path}: {str(e)}"


@tool
def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file (Microsoft Word document).

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Extracted text with paragraph information
    """
    try:
        doc = Document(docx_path)
        text = ""

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extract tables if any
        if doc.tables:
            text += "\n--- TABLES ---\n"
            for table_num, table in enumerate(doc.tables, 1):
                text += f"\nTable {table_num}:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"

        return text if text else "No text found in DOCX"
    except Exception as e:
        return f"Error reading DOCX {docx_path}: {str(e)}"
